import os
import re
import json
import uuid
import logging
import subprocess
from pathlib import Path
from copy import deepcopy

from flask import Flask, request, jsonify, send_file, abort, url_for
from flask_cors import CORS
from werkzeug.utils import secure_filename
from docx import Document
from dotenv import load_dotenv
from anthropic import Anthropic

load_dotenv()

# ---------- Config ----------
UPLOAD_DIR = Path(os.getenv("UPLOAD_FOLDER", "/tmp/uploads"))
RESULTS_DIR = Path(os.getenv("RESULTS_FOLDER", "/tmp/results"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
RESULTS_DIR.mkdir(parents=True, exist_ok=True)

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-6")
MAX_OUTPUT_TOKENS = int(os.getenv("CLAUDE_MAX_TOKENS", "8000"))

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("resume-enhancer")

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20MB
CORS(app)

_anthropic = Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None


# ---------- Health ----------
@app.get("/")
def health():
    return jsonify({
        "ok": True,
        "ai_configured": bool(_anthropic),
        "model": CLAUDE_MODEL,
    }), 200


@app.get("/ls")
def list_results():
    items = [{"name": p.name, "bytes": p.stat().st_size}
             for p in sorted(RESULTS_DIR.glob("*"))]
    return jsonify({"results": items}), 200


# ---------- DOCX helpers ----------
def extract_paragraphs(doc: Document):
    """Return list of {index, text} for non-empty paragraphs."""
    out = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            out.append({"index": i, "text": text})
    return out


def replace_paragraph_text(paragraph, new_text: str):
    """Replace paragraph text while preserving the first run's formatting."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    first_run = paragraph.runs[0]
    # Clear runs after the first
    for run in paragraph.runs[1:]:
        run.text = ""
    first_run.text = new_text


SYSTEM_PROMPT = """You are an expert resume editor. You rewrite resume bullets and summary text so the candidate's existing experience is presented in language that aligns with a target job description — without inventing skills, employers, dates, projects, or accomplishments the candidate has not actually done.

Hard rules:
- NEVER fabricate experience, technologies, employers, schools, dates, metrics, or job titles.
- Only rewrite content the candidate already has. You may reframe phrasing, surface relevant keywords from the job description that genuinely apply, and tighten language.
- Preserve all proper nouns: people, company names, school names, project names, dates, locations.
- Do NOT rewrite contact info, names, section headers (like "Experience", "Education", "Skills"), dates, or single-word labels.
- Keep each rewritten line roughly the same length as the original (±25%). Bullets stay bullets.
- Use strong action verbs and quantifiable outcomes ONLY where they were present in the original. Do not invent numbers.
- Match tense to the original (past roles → past tense; current role → present tense).
- Output strictly valid JSON, no prose, no markdown fences.
"""

USER_PROMPT_TEMPLATE = """Job description:
<job_description>
{jd}
</job_description>

Resume paragraphs (indexed). Some are headers/names/contact lines — leave those unchanged by omitting them from the response:
<resume_paragraphs>
{paragraphs_json}
</resume_paragraphs>

Return JSON in exactly this shape:
{{
  "changes": [
    {{"index": <int from input>, "new_text": "<rewritten paragraph>"}}
  ]
}}

Only include paragraphs you actually rewrite. Skip section headers, the candidate's name, contact info, dates, and short labels. Aim to rewrite 60-90% of the substantive bullets and the summary if present."""


def call_claude(paragraphs, jd: str) -> dict:
    """Call Claude and return parsed JSON {changes: [{index, new_text}]}."""
    if not _anthropic:
        raise RuntimeError("ANTHROPIC_API_KEY is not configured on the server.")

    user_content = USER_PROMPT_TEMPLATE.format(
        jd=jd,
        paragraphs_json=json.dumps(paragraphs, ensure_ascii=False, indent=2),
    )

    log.info("Calling Claude (%s) with %d paragraphs", CLAUDE_MODEL, len(paragraphs))
    resp = _anthropic.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=MAX_OUTPUT_TOKENS,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_content}],
    )

    raw = "".join(block.text for block in resp.content if getattr(block, "type", "") == "text")
    raw = raw.strip()

    # Strip accidental markdown fences
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        log.error("Claude returned non-JSON: %s", raw[:500])
        raise RuntimeError(f"AI response was not valid JSON: {e}")

    changes = data.get("changes", [])
    if not isinstance(changes, list):
        raise RuntimeError("AI response missing 'changes' array.")

    return {"changes": changes}


def apply_changes(doc: Document, changes) -> int:
    """Apply {index, new_text} changes to the doc. Returns count applied."""
    by_index = {c["index"]: c["new_text"] for c in changes
                if isinstance(c, dict) and "index" in c and "new_text" in c}
    applied = 0
    for i, para in enumerate(doc.paragraphs):
        if i in by_index:
            new_text = by_index[i].strip()
            if new_text:
                replace_paragraph_text(para, new_text)
                applied += 1
    return applied


def convert_to_pdf(src_docx: Path, out_dir: Path):
    """Convert DOCX -> PDF using LibreOffice; return PDF path or None."""
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf",
               "--outdir", str(out_dir), str(src_docx)]
        log.info("Running: %s", " ".join(cmd))
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if res.returncode != 0:
            log.warning("soffice exit=%s stderr=%s", res.returncode, res.stderr[:400])
            return None
        pdf_path = out_dir / (src_docx.stem + ".pdf")
        return pdf_path if pdf_path.exists() else None
    except FileNotFoundError:
        log.warning("LibreOffice not installed; skipping PDF preview.")
        return None
    except Exception as e:
        log.warning("PDF conversion failed: %s", e)
        return None


# ---------- Routes ----------
@app.route("/enhance", methods=["POST", "OPTIONS"])
def enhance():
    if request.method == "OPTIONS":
        return ("", 204)

    file = (request.files.get("resume")
            or request.files.get("file")
            or request.files.get("upload"))
    jd = (request.form.get("jobDescription")
          or request.form.get("job_description")
          or request.form.get("description")
          or "").strip()

    if not file or not file.filename:
        return jsonify({"status": "error", "message": "Missing 'resume' file."}), 400
    if not file.filename.lower().endswith(".docx"):
        return jsonify({"status": "error", "message": "Only .docx resumes are supported."}), 400
    if not jd:
        return jsonify({"status": "error", "message": "Missing 'jobDescription'."}), 400
    if not _anthropic:
        return jsonify({
            "status": "error",
            "message": "Server is missing ANTHROPIC_API_KEY. Set it in the environment."
        }), 503

    # Save upload
    in_name = secure_filename(f"{uuid.uuid4()}.docx")
    in_path = UPLOAD_DIR / in_name
    file.save(in_path)
    log.info("Upload saved -> %s (%d bytes)", in_path, in_path.stat().st_size)

    # Open & extract
    try:
        doc = Document(str(in_path))
    except Exception as e:
        log.exception("Failed to open DOCX")
        return jsonify({"status": "error", "message": f"Could not open DOCX: {e}"}), 400

    paragraphs = extract_paragraphs(doc)
    if not paragraphs:
        return jsonify({"status": "error", "message": "Resume appears to be empty."}), 400

    # Call Claude
    try:
        result = call_claude(paragraphs, jd)
    except Exception as e:
        log.exception("Claude call failed")
        return jsonify({"status": "error", "message": f"AI enhancement failed: {e}"}), 502

    applied = apply_changes(doc, result["changes"])
    log.info("Applied %d/%d changes", applied, len(result["changes"]))

    # Save enhanced DOCX
    out_name = secure_filename(f"{uuid.uuid4()}_enhanced.docx")
    out_path = RESULTS_DIR / out_name
    doc.save(str(out_path))

    pdf_path = convert_to_pdf(out_path, RESULTS_DIR)
    pdf_url = url_for("serve_result", fname=pdf_path.name, _external=True) if pdf_path else None
    docx_url = url_for("serve_result", fname=out_name, _external=True)

    return jsonify({
        "status": "success",
        "docx_url": docx_url,
        "pdf_url": pdf_url,
        "changes_applied": applied,
        "paragraphs_total": len(paragraphs),
        "message": None if pdf_url else "PDF preview unavailable (LibreOffice not installed).",
    }), 200


@app.get("/results/<path:fname>")
def serve_result(fname):
    fp = RESULTS_DIR / fname
    if not fp.exists():
        abort(404)
    if fp.suffix.lower() == ".docx":
        return send_file(
            fp,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=fname,
        )
    if fp.suffix.lower() == ".pdf":
        return send_file(fp, mimetype="application/pdf", as_attachment=False)
    return send_file(fp, as_attachment=True, download_name=fname)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "10000"))
    app.run(host="0.0.0.0", port=port, debug=os.getenv("FLASK_DEBUG") == "1")
